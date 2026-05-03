"""Format-preserving resume rewrite via the agent graph."""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from api.agents import run_agent

log = logging.getLogger(__name__)

@dataclass
class _TextNode:
    obj: Paragraph | _Cell
    text: str
    is_header: bool


def extract_lines(docx_bytes: bytes) -> list[str]:
    """Public: extract non-empty text lines from a DOCX file. Used for diff views."""
    from io import BytesIO
    doc = Document(BytesIO(docx_bytes))
    return [n.text for n in _extract_nodes(doc)]


def _extract_nodes(doc: Document) -> list[_TextNode]:
    """Walk paragraphs and table cells, return all non-empty text nodes."""
    nodes: list[_TextNode] = []

    def add(obj: Paragraph | _Cell, text: str) -> None:
        t = text.strip()
        if t:
            nodes.append(_TextNode(obj=obj, text=t, is_header=False))

    for para in doc.paragraphs:
        add(para, para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    add(para, para.text)

    return nodes


def _set_paragraph_text(para: Paragraph, new_text: str) -> None:
    """Replace all runs in a paragraph with new_text, preserving first run's formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    # Keep first run's format, clear others
    first = para.runs[0]
    first.text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _set_cell_text(cell: _Cell, original: str, new_text: str) -> None:
    """Replace matching paragraph text in a cell."""
    for para in cell.paragraphs:
        if para.text.strip() == original:
            _set_paragraph_text(para, new_text)
            return


async def rewrite(
    docx_bytes: bytes,
    job_description: str,
    rag_tag: str,
) -> bytes:
    """
    Rewrite resume bullet points to match `job_description`.
    Returns new DOCX bytes.
    """
    doc = Document(io.BytesIO(docx_bytes))
    nodes = _extract_nodes(doc)

    if not nodes:
        log.warning("resume_engine: no text nodes found in DOCX")
        return docx_bytes

    # Build numbered bullet list for the agent
    numbered = "\n".join(f"{i + 1}. {n.text}" for i, n in enumerate(nodes))

    user_prompt = (
        f"<job_description>\n{job_description[:3000]}\n</job_description>\n\n"
        f"<resume_bullets>\n{numbered}\n</resume_bullets>\n\n"
        "Rewrite the bullets to match this role. Return the same numbered list."
    )

    final = await run_agent(
        mode="rewrite_resume",
        user_prompt=user_prompt,
        rag_tag=rag_tag,
        job_description=job_description,
    )
    response = final.get("output", "")
    rewrites = _parse_numbered(response, expected=len(nodes))

    # Apply rewrites back into the document
    for i, node in enumerate(nodes):
        new_text = rewrites.get(i, node.text)
        if isinstance(node.obj, Paragraph):
            _set_paragraph_text(node.obj, new_text)
        else:
            _set_cell_text(node.obj, node.text, new_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_numbered(response: str, expected: int) -> dict[int, str]:
    """Parse '1. text' lines from Claude response into 0-indexed dict."""
    out: dict[int, str] = {}
    for line in response.splitlines():
        line = line.strip()
        if not line:
            continue
        dot = line.find(". ")
        if dot == -1:
            continue
        num_part = line[:dot]
        if not num_part.isdigit():
            continue
        idx = int(num_part) - 1
        if 0 <= idx < expected:
            out[idx] = line[dot + 2:].strip()
    return out
